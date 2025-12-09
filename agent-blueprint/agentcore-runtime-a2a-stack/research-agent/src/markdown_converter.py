"""
Markdown to DOCX Converter

Converts markdown content to Word document with:
- Proper heading hierarchy
- Inline formatting (bold, italic, code)
- Lists (bullet and numbered)
- Images/charts embedding
- URL citations as footnotes
"""

import re
import os
import logging
from typing import List, Dict, Optional
from pathlib import Path

from docx import Document
from docx.shared import Pt, Inches
from docx.oxml.shared import OxmlElement
from docx.oxml.ns import qn

logger = logging.getLogger(__name__)


def parse_inline_formatting(text: str, paragraph, doc=None):
    """
    Parse inline markdown formatting and add formatted runs to paragraph.

    Supports:
    - **bold**
    - *italic*
    - `code`
    - [1], [2] citations (superscript)
    - §FOOTNOTE:number:url§ markers (converted to superscript numbers)

    Args:
        text: Text with inline markdown
        paragraph: python-docx paragraph object
        doc: Document object (for tracking footnotes)
    """
    patterns = [
        (r'\*\*(.+?)\*\*', 'bold'),
        (r'\*(.+?)\*', 'italic'),
        (r'`(.+?)`', 'code'),
        (r'§FOOTNOTE:(\d+):([^§]+)§', 'footnote'),
        (r'\[\d+(?:,\s*\d+)*\]', 'citation')
    ]

    pos = 0
    while pos < len(text):
        earliest_match = None
        earliest_pos = len(text)
        match_type = None

        for pattern, fmt_type in patterns:
            match = re.search(pattern, text[pos:])
            if match and match.start() < earliest_pos - pos:
                earliest_match = match
                earliest_pos = pos + match.start()
                match_type = fmt_type

        if not earliest_match:
            if pos < len(text):
                paragraph.add_run(text[pos:])
            break

        if earliest_pos > pos:
            paragraph.add_run(text[pos:earliest_pos])

        matched_text = earliest_match.group(0)

        if match_type == 'bold':
            run = paragraph.add_run(earliest_match.group(1))
            run.bold = True
        elif match_type == 'italic':
            run = paragraph.add_run(earliest_match.group(1))
            run.italic = True
        elif match_type == 'code':
            run = paragraph.add_run(earliest_match.group(1))
            run.font.name = 'Courier New'
            run.font.size = Pt(10)
        elif match_type == 'footnote':
            footnote_num = earliest_match.group(1)
            footnote_url = earliest_match.group(2)
            run = paragraph.add_run(footnote_num)
            run.font.superscript = True
            run.font.size = Pt(9)
            if doc and hasattr(doc, '_footnote_map'):
                doc._footnote_map[footnote_num] = footnote_url
        elif match_type == 'citation':
            run = paragraph.add_run(matched_text)
            run.font.superscript = True
            run.font.size = Pt(9)

        pos = earliest_pos + len(matched_text)


def add_hyperlink(paragraph, url: str, text: str):
    """Add a hyperlink to a paragraph."""
    part = paragraph.part
    r_id = part.relate_to(
        url,
        'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink',
        is_external=True
    )

    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)

    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')

    color = OxmlElement('w:color')
    color.set(qn('w:val'), '0563C1')
    rPr.append(color)

    u = OxmlElement('w:u')
    u.set(qn('w:val'), 'single')
    rPr.append(u)

    new_run.append(rPr)

    text_elem = OxmlElement('w:t')
    text_elem.text = text
    new_run.append(text_elem)

    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


def markdown_to_docx(
    markdown_content: str,
    output_path: str,
    chart_files: Optional[List[Dict[str, str]]] = None
) -> str:
    """
    Convert markdown to Word document.

    Args:
        markdown_content: Markdown content
        output_path: Output .docx file path
        chart_files: List of chart dicts with 'path', 'title'

    Returns:
        Path to created document
    """
    # Build URL citation mapping
    url_pattern = r'\[(https?://[^\]]+)\]'
    urls_found = re.findall(url_pattern, markdown_content)

    url_to_number = {}
    citation_counter = 1
    for url in urls_found:
        if url not in url_to_number:
            url_to_number[url] = citation_counter
            citation_counter += 1

    # Replace URLs with footnote markers
    def replace_url_with_marker(match):
        url = match.group(1)
        number = url_to_number.get(url, '?')
        return f'§FOOTNOTE:{number}:{url}§'

    markdown_content = re.sub(url_pattern, replace_url_with_marker, markdown_content)

    doc = Document()
    doc._footnote_map = {}

    lines = markdown_content.split('\n')
    charts_inserted = set()
    skip_next = False

    for i, line in enumerate(lines):
        line = line.rstrip()

        if not line:
            continue

        if skip_next:
            skip_next = False
            continue

        # Heading level 1
        if line.startswith('# '):
            heading_text = line[2:]
            doc.add_heading(heading_text, level=0)

        # Heading level 2
        elif line.startswith('## '):
            heading_text = line[3:]
            doc.add_heading(heading_text, level=1)

            # Insert charts after Executive Summary
            if chart_files and heading_text.strip().lower() == 'executive summary':
                for chart_info in chart_files:
                    chart_id = chart_info.get('id', '')
                    if chart_id in charts_inserted:
                        continue

                    chart_path = chart_info.get('path')
                    chart_title = chart_info.get('title', chart_id)

                    if chart_path and os.path.exists(chart_path):
                        try:
                            doc.add_paragraph()
                            title_para = doc.add_paragraph()
                            title_run = title_para.add_run(chart_title)
                            title_run.bold = True

                            doc.add_picture(chart_path, width=Inches(6))
                            doc.add_paragraph()
                            charts_inserted.add(chart_id)
                            logger.info(f"Inserted chart: {chart_title}")
                        except Exception as e:
                            logger.warning(f"Failed to insert chart {chart_title}: {e}")

        # Heading level 3
        elif line.startswith('### '):
            heading_text = line[4:]
            doc.add_heading(heading_text, level=2)

        # Heading level 4
        elif line.startswith('#### '):
            heading_text = line[5:]
            doc.add_heading(heading_text, level=3)

        # Image (markdown syntax)
        elif line.startswith('!['):
            match = re.match(r'!\[(.*?)\]\((.*?)\)', line)
            if match:
                alt_text, image_path = match.groups()

                if os.path.exists(image_path):
                    try:
                        doc.add_picture(image_path, width=Inches(6))

                        # Check for figure caption
                        if i + 1 < len(lines):
                            next_line = lines[i + 1].strip()
                            if next_line.startswith('*Figure ') and next_line.endswith('*'):
                                caption_text = next_line[1:-1]
                                p = doc.add_paragraph()
                                run = p.add_run(caption_text)
                                run.italic = True
                                p.alignment = 1  # Center
                                skip_next = True
                    except Exception as e:
                        logger.warning(f"Failed to insert image {image_path}: {e}")
                        p = doc.add_paragraph()
                        run = p.add_run(f"[Image: {alt_text}]")
                        run.italic = True
                else:
                    p = doc.add_paragraph()
                    run = p.add_run(f"[Image not found: {alt_text}]")
                    run.italic = True

        # Horizontal rule
        elif line.startswith('---'):
            doc.add_paragraph('_' * 50)

        # Bullet list
        elif line.startswith('- ') or line.startswith('* '):
            p = doc.add_paragraph(style='List Bullet')
            parse_inline_formatting(line[2:], p, doc)

        # Numbered list
        elif re.match(r'^\d+\.\s', line):
            text = re.sub(r'^\d+\.\s', '', line)
            p = doc.add_paragraph(style='List Number')
            parse_inline_formatting(text, p, doc)

        # Italic line (metadata)
        elif line.startswith('*') and line.endswith('*') and not line.startswith('**'):
            p = doc.add_paragraph()
            run = p.add_run(line[1:-1])
            run.italic = True

        # Regular paragraph
        else:
            p = doc.add_paragraph()
            parse_inline_formatting(line, p, doc)

    # Add Footnotes section
    if hasattr(doc, '_footnote_map') and doc._footnote_map:
        doc.add_paragraph()
        doc.add_page_break()
        doc.add_heading('Footnotes', level=2)

        sorted_footnotes = sorted(doc._footnote_map.items(), key=lambda x: int(x[0]))

        for number, url in sorted_footnotes:
            p = doc.add_paragraph()
            run = p.add_run(number)
            run.font.superscript = True
            run.font.size = Pt(9)
            p.add_run(' ')
            add_hyperlink(p, url, url)

    doc.save(output_path)
    logger.info(f"Document saved: {output_path}")

    return output_path


def docx_to_pdf(docx_path: str, pdf_path: str) -> str:
    """
    Convert Word document to PDF.

    Args:
        docx_path: Input .docx file path
        pdf_path: Output .pdf file path

    Returns:
        Path to created PDF
    """
    import platform
    import subprocess

    if not os.path.exists(docx_path):
        raise FileNotFoundError(f"DOCX file not found: {docx_path}")

    output_dir = os.path.dirname(pdf_path)
    os.makedirs(output_dir, exist_ok=True)

    if platform.system() == 'Linux':
        # Use LibreOffice on Linux
        result = subprocess.run(
            [
                'soffice',
                '--headless',
                '--convert-to', 'pdf',
                '--outdir', output_dir,
                docx_path
            ],
            capture_output=True,
            text=True,
            timeout=60
        )

        if result.returncode != 0:
            raise Exception(f"LibreOffice conversion failed: {result.stderr}")

        # Rename if needed
        libreoffice_output = os.path.join(
            output_dir,
            os.path.basename(docx_path).replace('.docx', '.pdf')
        )
        if libreoffice_output != pdf_path and os.path.exists(libreoffice_output):
            os.rename(libreoffice_output, pdf_path)
    else:
        # Use docx2pdf on Mac/Windows
        try:
            from docx2pdf import convert
            convert(docx_path, pdf_path)
        except ImportError:
            raise Exception("docx2pdf not installed. Install with: pip install docx2pdf")

    if not os.path.exists(pdf_path):
        raise Exception(f"PDF file was not created at {pdf_path}")

    logger.info(f"PDF created: {pdf_path}")
    return pdf_path
