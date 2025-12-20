"""
Word Document Tools - 5 essential tools for Word document management.

Tools:
1. store_word_document - Save uploaded Word files to workspace
2. create_word_document - Create new Word document from Markdown
3. modify_word_document - Modify existing Word document with python-docx code
4. list_my_word_documents - List all Word documents in workspace
5. get_word_document - Retrieve document for download

Pattern follows diagram_tool for Code Interpreter usage.
"""

import os
import re
import logging
from typing import Dict, Any, Optional
from strands import tool, ToolContext
from bedrock_agentcore.tools.code_interpreter_client import CodeInterpreter
from .lib.document_manager import WordDocumentManager

logger = logging.getLogger(__name__)


def _validate_document_name(name: str) -> tuple[bool, Optional[str]]:
    """Validate document name meets requirements (without extension).

    Rules:
    - Only letters (a-z, A-Z), numbers (0-9), and hyphens (-)
    - No spaces, underscores, or special characters
    - No consecutive hyphens
    - No leading/trailing hyphens

    Args:
        name: Document name without extension (e.g., "sales-report")

    Returns:
        (is_valid, error_message)
        - (True, None) if valid
        - (False, error_message) if invalid
    """
    # Check for empty name
    if not name:
        return False, "Document name cannot be empty"

    # Check for valid characters: only letters, numbers, hyphens
    if not re.match(r'^[a-zA-Z0-9\-]+$', name):
        invalid_chars = re.findall(r'[^a-zA-Z0-9\-]', name)
        return False, f"Invalid characters in name: {set(invalid_chars)}. Use only letters, numbers, and hyphens (-)."

    # Check for consecutive hyphens
    if '--' in name:
        return False, "Name cannot contain consecutive hyphens (--)"

    # Check for leading/trailing hyphens
    if name.startswith('-') or name.endswith('-'):
        return False, "Name cannot start or end with a hyphen"

    return True, None


def _sanitize_document_name_for_bedrock(filename: str) -> str:
    """Sanitize existing filename for Bedrock API (removes extension).

    Use this ONLY for existing files being read from S3.
    For new files, use _validate_document_name() instead.

    Args:
        filename: Original filename with extension (e.g., "test_document_v2.docx")

    Returns:
        Sanitized name without extension (e.g., "test-document-v2")
    """
    # Remove extension
    if '.' in filename:
        name, ext = filename.rsplit('.', 1)
    else:
        name = filename

    # Replace underscores and spaces with hyphens
    name = name.replace('_', '-').replace(' ', '-')

    # Keep only alphanumeric and hyphens
    name = re.sub(r'[^a-zA-Z0-9\-]', '', name)

    # Replace multiple consecutive hyphens with single hyphen
    name = re.sub(r'\-+', '-', name)

    # Trim hyphens from start/end
    name = name.strip('-')

    # If name becomes empty, use default
    if not name:
        name = 'document'

    if name != filename.replace('.docx', ''):
        logger.info(f"Sanitized document name for Bedrock: '{filename}' → '{name}'")

    return name


def _get_code_interpreter_id() -> Optional[str]:
    """Get Custom Code Interpreter ID from environment or Parameter Store"""
    # 1. Check environment variable (set by AgentCore Runtime)
    code_interpreter_id = os.getenv('CODE_INTERPRETER_ID')
    if code_interpreter_id:
        logger.info(f"Found CODE_INTERPRETER_ID in environment: {code_interpreter_id}")
        return code_interpreter_id

    # 2. Try Parameter Store (for local development or alternative configuration)
    try:
        import boto3
        project_name = os.getenv('PROJECT_NAME', 'strands-agent-chatbot')
        environment = os.getenv('ENVIRONMENT', 'dev')
        region = os.getenv('AWS_REGION', 'us-west-2')
        param_name = f"/{project_name}/{environment}/agentcore/code-interpreter-id"

        logger.info(f"Checking Parameter Store for Code Interpreter ID: {param_name}")
        ssm = boto3.client('ssm', region_name=region)
        response = ssm.get_parameter(Name=param_name)
        code_interpreter_id = response['Parameter']['Value']
        logger.info(f"Found CODE_INTERPRETER_ID in Parameter Store: {code_interpreter_id}")
        return code_interpreter_id
    except Exception as e:
        logger.warning(f"Custom Code Interpreter ID not found in Parameter Store: {e}")
        return None


def _get_user_session_ids(tool_context: ToolContext) -> tuple[str, str]:
    """Extract user_id and session_id from ToolContext

    Returns:
        (user_id, session_id) tuple
    """
    # Extract from invocation_state (set by agent)
    invocation_state = tool_context.invocation_state
    user_id = invocation_state.get('user_id', 'default_user')
    session_id = invocation_state.get('session_id', 'default_session')

    logger.info(f"Extracted IDs: user_id={user_id}, session_id={session_id}")
    return user_id, session_id


def _extract_uploaded_docx_files(tool_context: ToolContext) -> list[tuple[str, bytes]]:
    """Extract uploaded .docx files from ToolContext

    Returns:
        List of (filename, file_bytes) tuples for .docx files
    """
    uploaded_files = []

    # Check if files were uploaded in this invocation
    invocation_state = tool_context.invocation_state
    files = invocation_state.get('uploaded_files', [])

    for file_info in files:
        filename = file_info.get('filename', '')
        if filename.endswith('.docx'):
            file_bytes = file_info.get('bytes')
            if file_bytes:
                uploaded_files.append((filename, file_bytes))
                logger.info(f"Found uploaded .docx file: {filename} ({len(file_bytes)} bytes)")

    return uploaded_files


def _extract_uploaded_image_files(tool_context: ToolContext) -> list[tuple[str, bytes]]:
    """Extract uploaded image files from ToolContext

    Returns:
        List of (filename, file_bytes) tuples for image files (.png, .jpg, .jpeg, .gif)
    """
    uploaded_images = []

    # Check if files were uploaded in this invocation
    invocation_state = tool_context.invocation_state
    files = invocation_state.get('uploaded_files', [])

    for file_info in files:
        filename = file_info.get('filename', '')
        if filename.lower().endswith(('.png', '.jpg', '.jpeg', '.gif', '.bmp')):
            file_bytes = file_info.get('bytes')
            if file_bytes:
                uploaded_images.append((filename, file_bytes))
                logger.info(f"Found uploaded image file: {filename} ({len(file_bytes)} bytes)")

    return uploaded_images


def _upload_images_to_ci(code_interpreter: CodeInterpreter, images: list[tuple[str, bytes]]) -> list[str]:
    """Upload multiple image files to Code Interpreter workspace

    Args:
        code_interpreter: Active CodeInterpreter instance
        images: List of (filename, file_bytes) tuples

    Returns:
        List of uploaded filenames
    """
    uploaded = []

    for filename, file_bytes in images:
        try:
            import base64
            encoded_bytes = base64.b64encode(file_bytes).decode('utf-8')

            write_code = f"""
import base64

# Decode and write image file
file_bytes = base64.b64decode('{encoded_bytes}')
with open('{filename}', 'wb') as f:
    f.write(file_bytes)

print(f"Image uploaded: {filename} ({{len(file_bytes)}} bytes)")
"""

            response = code_interpreter.invoke("executeCode", {
                "code": write_code,
                "language": "python",
                "clearContext": False
            })

            # Check for errors
            for event in response.get("stream", []):
                result = event.get("result", {})
                if result.get("isError", False):
                    error_msg = result.get("structuredContent", {}).get("stderr", "Unknown error")
                    logger.error(f"Failed to upload image {filename}: {error_msg[:200]}")
                    continue

            uploaded.append(filename)
            logger.info(f"✅ Uploaded image to Code Interpreter: {filename}")

        except Exception as e:
            logger.error(f"Failed to upload image {filename}: {e}")
            continue

    return uploaded


@tool(context=True)
def store_word_document(
    tool_context: ToolContext,
    custom_filename: Optional[str] = None
) -> Dict[str, Any]:
    """Save uploaded Word document(s) to your workspace for editing and persistence.

    This tool detects .docx files uploaded in the conversation and saves them to:
    1. Code Interpreter session (for immediate editing)
    2. S3 storage (for persistence across sessions)

    IMPORTANT: When user uploads .docx file, call this tool first (uploaded files only available in current turn).

    Use this tool when:
    - User uploads a Word file
    - User says "save this file", "store this document"

    Args:
        custom_filename: Optional. Rename the file when saving (single file only).
                        Must end with .docx
                        Example: "quarterly_report.docx"

                        Leave empty to use original filename.

    Returns:
        Success message with:
        - List of saved files with sizes
        - Current workspace file list (for context)
        - Metadata for frontend download buttons

    Example Usage:
        # Save uploaded file
        User: [uploads report.docx] "Edit this"
        AI: store_word_document()
        AI: modify_word_document(...)

        # Rename while saving
        User: [uploads old.docx] "Save as 'new.docx'"
        AI: store_word_document(custom_filename="new.docx")

    Note:
        - Automatically detects all .docx files from uploaded files
        - Single file can be renamed with custom_filename
        - Files available for modification after saving
    """
    try:
        logger.info("=== store_word_document called ===")

        # Get user and session IDs
        user_id, session_id = _get_user_session_ids(tool_context)

        # Initialize document manager
        doc_manager = WordDocumentManager(user_id, session_id)

        # Extract uploaded .docx files
        uploaded_files = _extract_uploaded_docx_files(tool_context)

        if not uploaded_files:
            return {
                "content": [{
                    "text": "❌ **No Word documents found**\n\nPlease upload one or more .docx files to save them to your workspace."
                }],
                "status": "error"
            }

        # Validate custom_filename if provided
        if custom_filename:
            if len(uploaded_files) > 1:
                return {
                    "content": [{
                        "text": f"❌ **Cannot rename multiple files**\n\nYou uploaded {len(uploaded_files)} files, but custom_filename can only be used with a single file."
                    }],
                    "status": "error"
                }

            # Validate .docx extension
            if not custom_filename.endswith('.docx'):
                return {
                    "content": [{
                        "text": f"❌ **Invalid filename**\n\nFilename must end with .docx: {custom_filename}"
                    }],
                    "status": "error"
                }

        # Get Code Interpreter
        code_interpreter_id = _get_code_interpreter_id()
        if not code_interpreter_id:
            return {
                "content": [{
                    "text": "❌ **Code Interpreter not configured**\n\nCODE_INTERPRETER_ID not found in environment or Parameter Store."
                }],
                "status": "error"
            }

        region = os.getenv('AWS_REGION', 'us-west-2')
        code_interpreter = CodeInterpreter(region)
        code_interpreter.start(identifier=code_interpreter_id)

        try:
            saved_files = []

            # Save each file
            for original_filename, file_bytes in uploaded_files:
                # Use custom filename if provided (single file case)
                filename = custom_filename if custom_filename else original_filename

                # Sync to both S3 and Code Interpreter
                sync_result = doc_manager.sync_to_both(
                    code_interpreter,
                    filename,
                    file_bytes,
                    metadata={'original_filename': original_filename}
                )

                saved_files.append({
                    'filename': filename,
                    'size_kb': sync_result['s3_info']['size_kb'],
                    'original_filename': original_filename
                })

                logger.info(f"✅ Saved: {filename} ({sync_result['s3_info']['size_kb']})")

            # Get current workspace list
            workspace_docs = doc_manager.list_s3_documents()
            workspace_summary = doc_manager.format_file_list(workspace_docs)

            # Build response message
            if len(saved_files) == 1:
                file_info = saved_files[0]
                renamed = " (renamed)" if custom_filename else ""
                message = f"""✅ **Document saved successfully**{renamed}

**File**: {file_info['filename']} ({file_info['size_kb']})

{workspace_summary}"""
            else:
                files_list = "\n".join([f"  - {f['filename']} ({f['size_kb']})" for f in saved_files])
                message = f"""✅ **{len(saved_files)} documents saved successfully**

{files_list}

{workspace_summary}"""

            # Prepare metadata for frontend (download buttons)
            metadata = {
                "saved_files": [
                    {
                        "filename": f['filename'],
                        "s3_key": doc_manager.get_s3_key(f['filename']),
                        "size_kb": f['size_kb']
                    } for f in saved_files
                ]
            }

            return {
                "content": [{"text": message}],
                "status": "success",
                "metadata": metadata
            }

        finally:
            code_interpreter.stop()

    except Exception as e:
        logger.error(f"store_word_document failed: {e}")
        return {
            "content": [{
                "text": f"❌ **Failed to save document(s)**\n\n{str(e)}"
            }],
            "status": "error"
        }


@tool(context=True)
def create_word_document(
    python_code: str,
    document_name: str,
    tool_context: ToolContext
) -> Dict[str, Any]:
    """Create a new Word document using python-docx code.

    This tool executes python-docx code to create a document from scratch.
    Perfect for generating structured documents with headings, paragraphs, tables, and charts.

    Available libraries: python-docx, matplotlib, seaborn, pandas, numpy

    Use this tool when:
    - User asks to create/generate a new Word document
    - User wants a document with specific structure and formatting
    - User needs charts/diagrams in the initial document

    Args:
        python_code: Python code using python-docx to build the document.
                    The document is initialized as: doc = Document()
                    After your code, it's automatically saved.

                    DO NOT include Document() initialization or doc.save() calls.

                    IMPORTANT: Uploaded images are automatically available in Code Interpreter.
                    If user uploads image.png, you can directly use it: doc.add_picture('image.png')

                    Common Patterns:

                    Basic Structure:
                    ```python
doc.add_heading('Quarterly Report', level=1)
doc.add_heading('Executive Summary', level=2)
doc.add_paragraph('Revenue increased by 15%...')

# Table with data
table = doc.add_table(rows=4, cols=3)
table.style = 'Light Grid Accent 1'
table.rows[0].cells[0].text = 'Quarter'
table.rows[0].cells[1].text = 'Revenue'
                    ```

                    With Generated Chart:
                    ```python
import matplotlib.pyplot as plt
from docx.shared import Inches

doc.add_heading('Sales Analysis', level=1)

# Generate chart
plt.figure(figsize=(8, 5))
plt.bar(['Q1','Q2','Q3','Q4'], [100, 120, 150, 140])
plt.title('Quarterly Sales')
plt.savefig('sales.png', dpi=300, bbox_inches='tight')
plt.close()

# Insert chart
doc.add_paragraph().add_run().add_picture('sales.png', width=Inches(6))
doc.add_paragraph('Figure 1: Sales performance')
                    ```

                    With Uploaded Image:
                    ```python
from docx.shared import Inches

doc.add_heading('Product Catalog', level=1)
doc.add_paragraph('Our new product line:')

# User uploaded 'product.png' - directly use it
doc.add_paragraph().add_run().add_picture('product.png', width=Inches(5))
                    ```

                    With Hyperlinks:
                    ```python
from docx.oxml.shared import OxmlElement
from docx.oxml.ns import qn

def add_hyperlink(paragraph, text, url):
    part = paragraph.part
    r_id = part.relate_to(url, 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink', is_external=True)
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)
    new_run = OxmlElement('w:r')
    new_run.text = text
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)

para = doc.add_paragraph('Visit ')
add_hyperlink(para, 'our website', 'https://example.com')
                    ```

                    With Placeholders for Future Edits:
                    ```python
doc.add_heading('Analysis', level=2)
doc.add_paragraph('{{INSERT_CHART_HERE}}')  # Marker for modify_word_document
doc.add_paragraph('Summary text...')
                    ```

        document_name: Document name WITHOUT extension (.docx is added automatically)
                      Use ONLY letters, numbers, hyphens (no underscores or spaces)
                      Examples: "sales-report", "Q4-analysis", "report-final"

    Returns:
        Success message with file details and workspace list

    Note:
        - Document is saved to workspace for future editing with modify_word_document
        - Uploaded images are automatically available in Code Interpreter
        - Keep code focused on structure; use modify_word_document for complex refinements
    """
    try:
        logger.info("=== create_word_document called ===")
        logger.info(f"Document name: {document_name}")

        # Validate document name (without extension)
        is_valid, error_msg = _validate_document_name(document_name)
        if not is_valid:
            return {
                "content": [{
                    "text": f"❌ **Invalid document name**: {document_name}\n\n{error_msg}\n\n**Examples of valid names:**\n- sales-report\n- Q4-analysis\n- report-final-v2"
                }],
                "status": "error"
            }

        # Add .docx extension
        document_filename = f"{document_name}.docx"
        logger.info(f"Full filename: {document_filename}")

        # Get user and session IDs
        user_id, session_id = _get_user_session_ids(tool_context)

        # Initialize document manager
        doc_manager = WordDocumentManager(user_id, session_id)

        # Get Code Interpreter
        code_interpreter_id = _get_code_interpreter_id()
        if not code_interpreter_id:
            return {
                "content": [{
                    "text": "❌ **Code Interpreter not configured**\n\nCODE_INTERPRETER_ID not found in environment or Parameter Store."
                }],
                "status": "error"
            }

        region = os.getenv('AWS_REGION', 'us-west-2')
        code_interpreter = CodeInterpreter(region)
        code_interpreter.start(identifier=code_interpreter_id)

        try:
            # Extract and upload images if any
            uploaded_images = _extract_uploaded_image_files(tool_context)
            if uploaded_images:
                uploaded_filenames = _upload_images_to_ci(code_interpreter, uploaded_images)
                logger.info(f"Uploaded {len(uploaded_filenames)} images to Code Interpreter: {uploaded_filenames}")

            # Get Code Interpreter path for file (filename only, no subdirectory)
            ci_path = doc_manager.get_ci_path(document_filename)

            # Build document creation code
            creation_code = f"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Create new document
doc = Document()

# Execute user's creation code
{python_code}

# Save document
doc.save('{ci_path}')
print(f"Document created: {ci_path}")
"""

            # Execute creation
            response = code_interpreter.invoke("executeCode", {
                "code": creation_code,
                "language": "python",
                "clearContext": False
            })

            # Check for errors
            for event in response.get("stream", []):
                result = event.get("result", {})
                if result.get("isError", False):
                    error_msg = result.get("structuredContent", {}).get("stderr", "Unknown error")
                    logger.error(f"Creation failed: {error_msg[:500]}")
                    code_interpreter.stop()
                    return {
                        "content": [{
                            "text": f"❌ **Failed to create document**\n\n```\n{error_msg[:1000]}\n```\n\n💡 Check your python-docx code for syntax errors or incorrect API usage."
                        }],
                        "status": "error"
                    }

            logger.info("Document creation completed")

            # Download from Code Interpreter
            file_bytes = doc_manager.download_from_code_interpreter(code_interpreter, document_filename)

            # Save to S3 for persistence
            s3_info = doc_manager.save_to_s3(
                document_filename,
                file_bytes,
                metadata={'source': 'python_code_creation'}
            )

            # Get current workspace list
            workspace_docs = doc_manager.list_s3_documents()
            workspace_summary = doc_manager.format_file_list(workspace_docs)

            message = f"""✅ **Document created successfully**

**File**: {document_filename} ({s3_info['size_kb']})

{workspace_summary}"""

            # Return success message
            return {
                "content": [{"text": message}],
                "status": "success",
                "metadata": {
                    "filename": document_filename,
                    "tool_type": "word_document"
                }
            }

        finally:
            code_interpreter.stop()

    except Exception as e:
        logger.error(f"create_word_document failed: {e}")
        return {
            "content": [{
                "text": f"❌ **Failed to create document**\n\n{str(e)}"
            }],
            "status": "error"
        }


@tool(context=True)
def modify_word_document(
    source_name: str,
    output_name: str,
    python_code: str,
    tool_context: ToolContext
) -> Dict[str, Any]:
    """Modify existing Word document using python-docx code and save with a new name.

    This tool loads a document from workspace, executes python-docx code to modify it,
    and saves it with a new filename to preserve the original.

    Available libraries: python-docx, matplotlib, seaborn, pandas, numpy

    Use this tool when:
    - User wants to edit/modify/update an existing document
    - User asks to add content, charts, or images to a document
    - User wants to refine or change parts of a document

    IMPORTANT Safety Rules:
    - Always use different output_filename than source_filename (e.g., "report.docx" → "report_v2.docx")
    - Always check array lengths before accessing (len(doc.paragraphs))
    - Use try-except for operations that might fail

    Args:
        source_name: Document name to load (WITHOUT extension, must exist in workspace)
                    Example: "sales-report", "Q4-analysis"
        output_name: New document name (WITHOUT extension, must be different from source)
                    Use ONLY letters, numbers, hyphens (no underscores or spaces)
                    Example: "sales-report-v2", "Q4-analysis-final"
        python_code: Python code using python-docx library to modify document.
                    The document is loaded as: doc = Document('<filename>')
                    After modifications, it's automatically saved.

                    DO NOT include Document() initialization or doc.save() calls.

                    IMPORTANT: Uploaded images are automatically available in Code Interpreter.
                    If user uploads diagram.png, you can directly use it in your code.

                    Common Patterns:

                    Insert Chart at Marker:
                    ```python
import matplotlib.pyplot as plt
from docx.shared import Inches

# Generate chart
plt.figure(figsize=(8, 5))
plt.plot([1,2,3,4], [10, 20, 25, 30])
plt.title('Sales Trend')
plt.savefig('trend.png', dpi=300, bbox_inches='tight')
plt.close()

# Find marker and replace with chart
for para in doc.paragraphs:
    if '{{CHART}}' in para.text:
        para.clear()
        para.add_run().add_picture('trend.png', width=Inches(6))
        break
                    ```

                    Insert Uploaded Image at Specific Location:
                    ```python
from docx.shared import Inches

# User uploaded 'diagram.png' - insert after paragraph 5
if len(doc.paragraphs) > 5:
    para = doc.paragraphs[5]
    para.add_run().add_picture('diagram.png', width=Inches(6))
                    ```

                    Add Hyperlink:
                    ```python
from docx.oxml.shared import OxmlElement
from docx.oxml.ns import qn

def add_hyperlink(paragraph, text, url):
    part = paragraph.part
    r_id = part.relate_to(url, 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink', is_external=True)
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)
    new_run = OxmlElement('w:r')
    new_run.text = text
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)

# Add to end of document
para = doc.add_paragraph('For more info: ')
add_hyperlink(para, 'Click here', 'https://example.com')
                    ```

                    Preserve Formatting When Editing:
                    ```python
# Preserve existing formatting when modifying text
if len(doc.paragraphs) > 0:
    p = doc.paragraphs[0]
    if len(p.runs) > 0:
        # Copy original formatting
        original_run = p.runs[0]
        font_name = original_run.font.name
        font_size = original_run.font.size
        is_bold = original_run.font.bold

        # Clear and add new text with same formatting
        for run in p.runs:
            run.text = ''

        new_run = p.runs[0] if len(p.runs) > 0 else p.add_run()
        new_run.text = 'New text with preserved formatting'
        new_run.font.name = font_name
        new_run.font.size = font_size
        new_run.font.bold = is_bold
                    ```

    Returns:
        Success message with file details and workspace list

    Note:
        - Uploaded images are automatically available in Code Interpreter
        - Use 0-based indexing (first paragraph = index 0)
        - Document automatically synced to S3
    """
    try:
        logger.info("=== modify_word_document called ===")
        logger.info(f"Source: {source_name}, Output: {output_name}")

        # Validate output name format
        is_valid, error_msg = _validate_document_name(output_name)
        if not is_valid:
            return {
                "content": [{
                    "text": f"❌ **Invalid output name**: {output_name}\n\n{error_msg}\n\n**Examples of valid names:**\n- sales-report-v2\n- Q4-analysis-final\n- report-revised"
                }],
                "status": "error"
            }

        # Ensure source and output are different
        if source_name == output_name:
            return {
                "content": [{
                    "text": f"❌ **Invalid name**\n\nOutput name must be different from source name to preserve the original.\n\nSource: {source_name}\nOutput: {output_name}\n\n💡 Try: \"{source_name}-v2\""
                }],
                "status": "error"
            }

        # Add .docx extensions
        source_filename = f"{source_name}.docx"
        output_filename = f"{output_name}.docx"
        logger.info(f"Full filenames: {source_filename} → {output_filename}")

        # Get user and session IDs
        user_id, session_id = _get_user_session_ids(tool_context)

        # Initialize document manager
        doc_manager = WordDocumentManager(user_id, session_id)

        # Get Code Interpreter
        code_interpreter_id = _get_code_interpreter_id()
        if not code_interpreter_id:
            return {
                "content": [{
                    "text": "❌ **Code Interpreter not configured**\n\nCODE_INTERPRETER_ID not found in environment or Parameter Store."
                }],
                "status": "error"
            }

        region = os.getenv('AWS_REGION', 'us-west-2')
        code_interpreter = CodeInterpreter(region)
        code_interpreter.start(identifier=code_interpreter_id)

        try:
            # Extract and upload images if any
            uploaded_images = _extract_uploaded_image_files(tool_context)
            if uploaded_images:
                uploaded_filenames = _upload_images_to_ci(code_interpreter, uploaded_images)
                logger.info(f"Uploaded {len(uploaded_filenames)} images to Code Interpreter: {uploaded_filenames}")

            # Ensure source file is in Code Interpreter (load from S3 if needed)
            source_ci_path = doc_manager.ensure_file_in_ci(code_interpreter, source_filename)

            # Generate output path
            output_ci_path = doc_manager.get_ci_path(output_filename)

            # Build modification code
            modification_code = f"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Load source document
doc = Document('{source_ci_path}')

# Execute user's modification code
{python_code}

# Save to output file
doc.save('{output_ci_path}')
print(f"Document modified and saved: {output_ci_path}")
"""

            # Execute modification
            response = code_interpreter.invoke("executeCode", {
                "code": modification_code,
                "language": "python",
                "clearContext": False
            })

            # Check for errors
            for event in response.get("stream", []):
                result = event.get("result", {})
                if result.get("isError", False):
                    error_msg = result.get("structuredContent", {}).get("stderr", "Unknown error")
                    logger.error(f"Modification failed: {error_msg[:500]}")
                    code_interpreter.stop()
                    return {
                        "content": [{
                            "text": f"❌ **Modification failed**\n\n```\n{error_msg[:1000]}\n```\n\n💡 Check your python-docx code for syntax errors or incorrect API usage."
                        }],
                        "status": "error"
                    }

            logger.info("Document modification completed")

            # Download modified document from Code Interpreter
            file_bytes = doc_manager.download_from_code_interpreter(code_interpreter, output_filename)

            # Save to S3 with output filename
            s3_info = doc_manager.save_to_s3(
                output_filename,
                file_bytes,
                metadata={
                    'source': 'modification',
                    'source_filename': source_filename,
                    'modified_at': 'timestamp'
                }
            )

            # Get current workspace list
            workspace_docs = doc_manager.list_s3_documents()
            workspace_summary = doc_manager.format_file_list(workspace_docs)

            # Build success message
            message = f"""✅ **Document modified successfully**

**Source**: {source_filename}
**Saved as**: {output_filename} ({s3_info['size_kb']})

{workspace_summary}"""

            # Return success message with metadata for download button
            return {
                "content": [{"text": message}],
                "status": "success",
                "metadata": {
                    "filename": output_filename,
                    "tool_type": "word_document"
                }
            }

        finally:
            code_interpreter.stop()

    except FileNotFoundError as e:
        logger.error(f"Document not found: {e}")
        return {
            "content": [{
                "text": f"❌ **Document not found**: {source_filename}"
            }],
            "status": "error"
        }
    except Exception as e:
        logger.error(f"modify_word_document failed: {e}")
        return {
            "content": [{
                "text": f"❌ **Failed to modify document**\n\n{str(e)}"
            }],
            "status": "error"
        }


@tool(context=True)
def list_my_word_documents(
    tool_context: ToolContext
) -> Dict[str, Any]:
    """List all Word documents in workspace.

    Shows all .docx files in workspace with size and metadata.

    Use this tool when:
    - User asks "what Word files do I have?"
    - User says "show my documents", "list files"
    - Before modifying: verify document exists
    - User wants to see workspace contents

    No arguments needed.

    Returns:
        - Formatted list of all Word documents
        - Each entry shows: filename, size, last modified date
        - Total file count
        - Metadata for frontend download buttons

    Example Usage:
        Scenario 1 - Check available files:
            User: "What Word documents do I have?"
            AI: list_my_word_documents()
            → Shows: report.docx, proposal.docx, analysis.docx

        Scenario 2 - Before modifying:
            User: "Edit my report"
            AI: [Unclear which file]
            AI: list_my_word_documents()
            AI: "I found these documents: ... Which one should I modify?"

    Example Output:
        📁 Workspace (3 documents):
          - q4_report.docx (45.6 KB) - Modified: 2025-01-15
          - proposal.docx (32.1 KB) - Modified: 2025-01-14
          - analysis.docx (78.4 KB) - Modified: 2025-01-13

    Note:
        - Shows files from workspace
        - Empty workspace shows helpful message
        - Frontend renders download buttons automatically
    """
    try:
        logger.info("=== list_my_word_documents called ===")

        # Get user and session IDs
        user_id, session_id = _get_user_session_ids(tool_context)

        # Initialize document manager
        doc_manager = WordDocumentManager(user_id, session_id)

        # List documents from S3
        documents = doc_manager.list_s3_documents()

        # Format list
        workspace_summary = doc_manager.format_file_list(documents)

        if documents:
            message = workspace_summary
        else:
            message = workspace_summary

        # Prepare metadata for frontend (download buttons)
        metadata = {
            "documents": [
                {
                    "filename": doc['filename'],
                    "s3_key": doc['s3_key'],
                    "size_kb": doc['size_kb'],
                    "last_modified": doc['last_modified']
                } for doc in documents
            ]
        }

        return {
            "content": [{"text": message}],
            "status": "success",
            "metadata": metadata
        }

    except Exception as e:
        logger.error(f"list_my_word_documents failed: {e}")
        return {
            "content": [{
                "text": f"❌ **Failed to list documents**\n\n{str(e)}"
            }],
            "status": "error"
        }


@tool(context=True)
def read_word_document(
    document_name: str,
    tool_context: ToolContext
) -> Dict[str, Any]:
    """Read and retrieve a specific Word document.

    This tool loads a document from workspace and returns it as downloadable bytes.
    The document content is accessible to you (the agent) for analysis and answering questions.

    Use this tool when:
    - User asks about document contents: "What's in report.docx?", "Summarize this document"
    - User wants to analyze the document: "How many tables are in this file?", "What's the main topic?"
    - User explicitly requests download: "Send me [filename]", "I need [document]"
    - You need to verify document contents before modification

    IMPORTANT:
    - For creating new documents: use create_word_document
    - For modifying documents: use modify_word_document

    Args:
        document_name: Document name WITHOUT extension (.docx is added automatically)
                      Must exist in workspace.
                      Example: "report", "proposal", "Q4-analysis"

    Returns:
        - Document metadata (filename, size, S3 location)
        - Special metadata format for frontend download
        - Frontend automatically shows download button

    Example Usage:
        # Download request
        User: "Send me the report"
        AI: read_word_document("report.docx")

        # After creation
        User: "Create report and send it"
        AI: create_word_document(...)
        AI: read_word_document("report.docx")

    Note:
        - File must exist in workspace
        - Frontend handles download automatically
    """
    try:
        logger.info("=== read_word_document called ===")
        logger.info(f"Document name: {document_name}")

        # Add .docx extension
        document_filename = f"{document_name}.docx"
        logger.info(f"Full filename: {document_filename}")

        # Get user and session IDs
        user_id, session_id = _get_user_session_ids(tool_context)

        # Initialize document manager
        doc_manager = WordDocumentManager(user_id, session_id)

        # Load from S3
        file_bytes = doc_manager.load_from_s3(document_filename)

        # Get file info
        documents = doc_manager.list_s3_documents()
        doc_info = next((d for d in documents if d['filename'] == document_filename), None)

        if not doc_info:
            raise FileNotFoundError(f"Document not found: {document_filename}")

        message = f"""✅ **Document ready for download**

**File**: {document_filename} ({doc_info['size_kb']})
**Last Modified**: {doc_info['last_modified'].split('T')[0]}"""

        # Sanitize document name for Bedrock API (remove extension, handle legacy files)
        # This handles legacy files with underscores or spaces
        sanitized_name = _sanitize_document_name_for_bedrock(document_filename)

        # Return with downloadable bytes
        return {
            "content": [
                {"text": message},
                {
                    "document": {
                        "format": "docx",
                        "name": sanitized_name,
                        "source": {
                            "bytes": file_bytes
                        }
                    }
                }
            ],
            "status": "success",
            "metadata": {
                "filename": document_filename,
                "s3_key": doc_manager.get_s3_key(document_filename),
                "size_kb": doc_info['size_kb'],
                "last_modified": doc_info['last_modified']
            }
        }

    except FileNotFoundError as e:
        logger.error(f"Document not found: {e}")
        return {
            "content": [{
                "text": f"❌ **Document not found**: {document_filename}"
            }],
            "status": "error"
        }
    except Exception as e:
        logger.error(f"read_word_document failed: {e}")
        return {
            "content": [{
                "text": f"❌ **Failed to read document**\n\n{str(e)}"
            }],
            "status": "error"
        }
